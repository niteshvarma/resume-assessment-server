import pymupdf4llm
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import logging
import tempfile
import os
import re

# Configure logging
logger = logging.getLogger(__name__)

class PyMuPDFDocumentParser():
    def __init__(self) -> None:
        logger.debug("Inside PyMuPDFDocumentParser instance init")

    def parse_document(self, file_path):
        pdf_path = None
        original_file_path = file_path  # <=== Save the original uploaded file path

        try:
            logger.debug(f"Inside parse_document method. Parsing document: {file_path}")

            assert os.path.exists(file_path), f"File {file_path} not found."

            if file_path.endswith('.docx'):
                logger.debug("File is a .docx. Converting to PDF.")
                pdf_path = os.path.join(tempfile.gettempdir(), os.path.basename(file_path).replace('.docx', '.pdf'))

                self.convert_docx_to_pdf(file_path, pdf_path)
                logger.debug(f"Converted .docx to PDF. PDF saved at: {pdf_path}")

                # Switch to parsing the generated PDF
                file_path = pdf_path

            logger.debug(f"Loading data into PyMuPDF4LLM from {file_path}")
            docString = pymupdf4llm.to_markdown(f"{file_path}")
            docString = self.fix_mid_sentence_line_breaks(docString)

            logger.debug(f"Document parsed successfully.")
            return docString

        except Exception as e:
            logger.error(f"Error parsing document: {e}")
            return None

        finally:
            # Always attempt to remove temp files
            files_to_remove = []

            if original_file_path and os.path.exists(original_file_path):
                files_to_remove.append(original_file_path)
            if pdf_path and os.path.exists(pdf_path):
                files_to_remove.append(pdf_path)

            for temp_file in files_to_remove:
                try:
                    logger.debug(f"Removing temporary file: {temp_file}")
                    os.remove(temp_file)
                except Exception as e:
                    logger.error(f"Failed to remove temporary file {temp_file}: {e}")


    def _docx_to_pdf(self, document, pdf_path):
        """
        Converts a python-docx Document object to a PDF using reportlab.
        """
        c = canvas.Canvas(pdf_path)

        # Set default font and size
        c.setFont("Helvetica", 12)

        # Margin and line height settings
        x_margin = 50
        y_margin = 750
        line_height = 15

        # Write document paragraphs to PDF
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                c.drawString(x_margin, y_margin, text)
                y_margin -= line_height

                # Start a new page if the content exceeds the page height
                if y_margin < 50:
                    c.showPage()
                    y_margin = 750

        # Save the PDF
        c.save()
        logger.debug(f"Saved PDF to: {pdf_path}")


    def fix_mid_sentence_line_breaks(self, text):
        """
        Remove mid-sentence line breaks while preserving legitimate line breaks,
        such as bullet points, paragraph breaks, and headings.
        
        Args:
            text (str): The text to process.
        
        Returns:
            str: The processed text with mid-sentence line breaks removed.
        """
        # Step 1: Preserve bullet points and headings
        # Match patterns for bullet points, numbered lists, or headings (e.g., "1.", "Heading:")
        bullet_or_heading_pattern = r'^(\s*[-*•]|\d+\.\s|[A-Z].*:)' 
        
        # Step 2: Replace mid-sentence line breaks
        # Look for line breaks that do not follow punctuation (".", "!", "?") 
        # and are not part of a new bullet, heading, or paragraph
        cleaned_text = re.sub(
            r'(?<![.!?])\n(?!\s*[-*•]|\d+\.\s|[A-Z])', 
            ' ', 
            text
        )
        
        # Step 3: Remove extra spaces introduced during line break merging
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        return cleaned_text

    def convert_docx_to_pdf(self, docx_file_path, pdf_output_path):
        """
        Convert a .docx file to a PDF while preserving text structure and formatting.
        Args:
            docx_file_path (str): Path to the .docx file.
            pdf_output_path (str): Path to save the output PDF.
        """
        try:
            # Extract text from the .docx file
            extracted_text = self.extract_text_from_docx(docx_file_path)

            # Write the extracted text to a PDF
            self.write_text_to_pdf(pdf_output_path, extracted_text)

            print(f"Conversion successful! PDF saved at: {pdf_output_path}")
        except Exception as e:
            logger.debug(f"Error during conversion: {e}")

    def extract_text_from_docx(self, file_path):
        """
        Extract text from a .docx file, preserving paragraph breaks and bullet points.
        Args:
            file_path (str): Path to the .docx file.
        Returns:
            str: The extracted and cleaned text.
        """
        doc = Document(file_path)
        text_content = []

        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                # Preserve bullets and indentation
                if para.style.name.startswith('List'):
                    text_content.append(f"- {para.text.strip()}")
                else:
                    text_content.append(para.text.strip())
        
        # Join paragraphs with a double newline to indicate separation
        return "\n\n".join(text_content)
    
    def write_text_to_pdf(self, output_pdf_path, text_content):
        """
        Write text content to a PDF file, ensuring proper text wrapping.
        Args:
            output_pdf_path (str): Path to save the output PDF.
            text_content (str): Text content to write to the PDF.
        """
        # Create a PDF document
        pdf = SimpleDocTemplate(output_pdf_path)
        styles = getSampleStyleSheet()
        story = []

        # Split the content into paragraphs
        paragraphs = text_content.split("\n\n")
        for para in paragraphs:
            # Add each paragraph with the "Normal" style
            story.append(Paragraph(para, styles['Normal']))
            story.append(Spacer(1, 12))  # Add space between paragraphs

        # Build the PDF
        pdf.build(story)
